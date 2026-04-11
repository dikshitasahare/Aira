"""
AIRA File Processing Pipeline
Handles PDF, Image (OCR), and code files
"""
import os
import io
import base64
import logging
from pathlib import Path
from django.conf import settings

logger = logging.getLogger(__name__)


# ── PDF Processing ─────────────────────────────────────────────────────────────

def extract_pdf_text(file_data: bytes, file_name: str) -> dict:
    """
    Extract text from PDF using pdfplumber (primary) with pypdf fallback.
    Returns dict with text, page_count, method_used.
    """
    result = {
        'text': '',
        'page_count': 0,
        'method': '',
        'success': False,
        'error': ''
    }

    # ── Try pdfplumber first (better text extraction) ──
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(io.BytesIO(file_data)) as pdf:
            result['page_count'] = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_text.append(f"--- Page {i+1} of {result['page_count']} ---\n{text.strip()}")

        if pages_text:
            result['text'] = "\n\n".join(pages_text)
            result['method'] = 'pdfplumber'
            result['success'] = True
            return result
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # ── Fallback to pypdf ──
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_data))
        result['page_count'] = len(reader.pages)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(f"--- Page {i+1} of {result['page_count']} ---\n{text.strip()}")

        if pages_text:
            result['text'] = "\n\n".join(pages_text)
            result['method'] = 'pypdf'
            result['success'] = True
            return result
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pypdf failed: {e}")

    # ── Fallback to PyMuPDF ──
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_data, filetype="pdf")
        result['page_count'] = len(doc)
        pages_text = []
        for i, page in enumerate(doc):
            text = page.get_text()
            if text and text.strip():
                pages_text.append(f"--- Page {i+1} of {result['page_count']} ---\n{text.strip()}")
        doc.close()

        if pages_text:
            result['text'] = "\n\n".join(pages_text)
            result['method'] = 'pymupdf'
            result['success'] = True
            return result
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyMuPDF failed: {e}")

    result['error'] = (
        "Could not extract text from PDF. "
        "The PDF may be scanned/image-based. "
        "Try converting to image and uploading again."
    )
    return result


def chunk_text(text: str, chunk_size: int = 3000, overlap: int = 300) -> list:
    """
    Split large text into overlapping chunks for LLM processing.
    Returns list of chunk strings.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        # Try to break at paragraph boundary
        if end < len(text):
            break_point = text.rfind('\n\n', start, end)
            if break_point == -1:
                break_point = text.rfind('\n', start, end)
            if break_point == -1:
                break_point = text.rfind('. ', start, end)
            if break_point != -1:
                end = break_point + 1
        chunks.append(text[start:end].strip())
        start = end - overlap  # overlap for context continuity

    return [c for c in chunks if c.strip()]


def prepare_pdf_context(text: str, max_chars: int = 12000) -> str:
    """
    Prepare PDF text for LLM — chunk if needed, return best context.
    """
    if len(text) <= max_chars:
        return text

    # Take beginning + end (most important parts of a document)
    half = max_chars // 2
    beginning = text[:half]
    ending = text[-half:]
    middle_note = f"\n\n[... middle section truncated for length — {len(text)} total chars ...]\n\n"
    return beginning + middle_note + ending


# ── Image Processing ────────────────────────────────────────────────────────────

def extract_image_text_ocr(file_data: bytes) -> dict:
    """
    Extract text from image using Tesseract OCR.
    Returns dict with text, success, error.
    """
    result = {'text': '', 'success': False, 'error': ''}

    try:
        import pytesseract
        from PIL import Image

        # Set tesseract path from settings
        tesseract_path = getattr(settings, 'TESSERACT_PATH', None)
        if tesseract_path and os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        image = Image.open(io.BytesIO(file_data))

        # Convert to RGB if needed
        if image.mode not in ('RGB', 'L'):
            image = image.convert('RGB')

        # Extract text
        text = pytesseract.image_to_string(image, config='--psm 3')
        result['text'] = text.strip()
        result['success'] = True

    except ImportError:
        result['error'] = "pytesseract not installed"
    except Exception as e:
        result['error'] = str(e)

    return result


def image_to_base64(file_data: bytes, mime_type: str) -> str:
    """Convert image bytes to base64 string for vision API."""
    return base64.b64encode(file_data).decode('utf-8')


def get_image_info(file_data: bytes) -> dict:
    """Get basic image metadata."""
    info = {'width': 0, 'height': 0, 'mode': '', 'format': ''}
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(file_data))
        info['width']  = img.width
        info['height'] = img.height
        info['mode']   = img.mode
        info['format'] = img.format or ''
    except Exception:
        pass
    return info


# ── File Validation ─────────────────────────────────────────────────────────────

def validate_file(file) -> dict:
    """
    Validate uploaded file for security.
    Returns {'valid': bool, 'error': str}
    """
    max_size = getattr(settings, 'MAX_UPLOAD_SIZE', 20 * 1024 * 1024)
    allowed_extensions = getattr(settings, 'ALLOWED_UPLOAD_EXTENSIONS', [])

    # Check file size
    if file.size > max_size:
        size_mb = max_size / (1024 * 1024)
        return {'valid': False, 'error': f'File too large. Maximum size is {size_mb:.0f}MB.'}

    # Check extension
    ext = Path(file.name).suffix.lower()
    if allowed_extensions and ext not in allowed_extensions:
        return {'valid': False, 'error': f'File type "{ext}" not allowed.'}

    # Check for null bytes (basic malware detection)
    if b'\x00' in file.name.encode():
        return {'valid': False, 'error': 'Invalid filename.'}

    return {'valid': True, 'error': ''}


def detect_file_type(file) -> str:
    """Detect file type from mime type and extension."""
    mime = getattr(file, 'content_type', '')
    name = getattr(file, 'name', '').lower()

    if mime == 'application/pdf' or name.endswith('.pdf'):
        return 'pdf'
    if mime.startswith('image/'):
        return 'image'
    if name.endswith('.docx') or name.endswith('.doc') or \
       mime in ('application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document',):
        return 'docx'
    if any(name.endswith(ext) for ext in [
        '.py', '.js', '.ts', '.html', '.css', '.json', '.md',
        '.txt', '.csv', '.xml', '.yaml', '.yml', '.sql', '.sh',
        '.cpp', '.c', '.java', '.go', '.rs'
    ]):
        return 'code'
    if mime.startswith('text/'):
        return 'text'
    return 'other'
    
# ── DOCX Processing ────────────────────────────────────

def extract_docx_text(file_data: bytes, file_name: str) -> dict:
    """Extract text from Word .docx files"""
    result = {
        'text': '', 'success': False,
        'error': '', 'method': 'python-docx'
    }
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(file_data))

        full_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading styles
                if para.style.name.startswith('Heading'):
                    level = para.style.name.split()[-1]
                    full_text.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {para.text}\n")
                else:
                    full_text.append(para.text)

        # Extract tables
        for table in doc.tables:
            full_text.append("\n[TABLE]\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                full_text.append(' | '.join(cells))
            full_text.append("[/TABLE]\n")

        result['text']    = '\n'.join(full_text)
        result['success'] = bool(result['text'].strip())
        if not result['success']:
            result['error'] = 'No text found in document'

    except ImportError:
        result['error'] = 'python-docx not installed. Run: pip install python-docx'
    except Exception as e:
        result['error'] = str(e)

    return result